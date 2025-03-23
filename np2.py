import streamlit as st
from openai import OpenAI
import os
import io
from docx import Document
from docx.shared import Pt

# ========== CONFIGURAÇÃO DA API ==========
api_key_env = os.getenv("OPENAI_API_KEY")

st.set_page_config(page_title="NutriChef Pro", page_icon="🥦")
st.image("NIDUp.png", width=250, caption="Núcleo de Inteligência de Dados")

st.title("🥦 NutriChef Pro – Receita inteligente com IA")

if not api_key_env:
    st.warning("🔐 Sua chave da OpenAI não foi encontrada no ambiente.")
    api_key_input = st.text_input("Insira sua chave da OpenAI aqui para continuar:", type="password")
    if not api_key_input:
        st.stop()
    client = OpenAI(api_key=api_key_input)
else:
    client = OpenAI(api_key=api_key_env)

st.markdown("**Preencha o formulário para receber sua receita personalizada!**")

# ========== FORMULÁRIO ==========
with st.form("form_receita"):
    nome = st.text_input("Seu nome (opcional)", "")
    objetivo = st.selectbox("Objetivo principal:", ["Hipertrofia", "Cutting", "Endurance", "Manutenção"])
    fase = st.selectbox("Fase do treino:", ["Bulk", "Definição", "Manutenção", "Não sei"])
    dieta = st.selectbox("Estilo alimentar:", ["Tradicional", "Low-carb", "Vegana", "Vegetariana", "Outra"])
    restricoes = st.multiselect("Restrições alimentares:", ["Sem lactose", "Sem glúten", "Não consumo ovos", "Nenhuma"])

    equipamentos_lista = ["Fogão", "Micro-ondas", "Airfryer", "Liquidificador", "Forno", "Só preparo frio"]
    equipamentos_selecao = st.multiselect("Equipamentos disponíveis:", ["Todos"] + equipamentos_lista)

    if "Todos" in equipamentos_selecao:
        equipamentos_selecao = equipamentos_lista

    tempo = st.selectbox("Tempo disponível para cozinhar:", ["Até 10 minutos", "Até 20 minutos", "Até 30 minutos", "Tempo livre"])
    refeicao = st.selectbox("Tipo de refeição:", ["Café da manhã", "Almoço", "Pré-treino", "Pós-treino", "Jantar", "Aleatória"])
    preferencia = st.text_input("Algum alimento que ama ou quer evitar?", "")
    dica = st.radio("Deseja uma dica de performance com a receita?", ["Sim", "Não"])

    enviar = st.form_submit_button("🔍 Gerar Receita")

# ========== GERAÇÃO COM IA ==========
if enviar:
    restricoes_str = ", ".join(restricoes) if restricoes else "Nenhuma"
    equipamentos_str = ", ".join(equipamentos_selecao) if equipamentos_selecao else "Não informado"

    prompt = f"""
Você é um nutricionista esportivo especializado em alimentação de performance.
Gere uma receita simples e prática com base nas informações abaixo:

Usuário: {nome or 'Usuário'}
Objetivo: {objetivo}
Fase do treino: {fase}
Estilo alimentar: {dieta}
Restrições: {restricoes_str}
Equipamentos: {equipamentos_str}
Tempo disponível: {tempo}
Tipo de refeição: {refeicao}
Preferências: {preferencia or 'Nenhuma'}
Deseja dica de performance? {dica}

Formato da resposta:
1. Nome da receita
2. Objetivo declarado
3. Tempo de preparo
4. Ingredientes com quantidades
5. Modo de preparo passo a passo
6. Macros estimados (proteína, carboidrato, gordura, calorias)
7. Dica extra (se solicitado)

Dependendo do caso, se for pertinente, considere a tabela base a seguir :

Formato: Tabela base (pode ser implementado em CSV, Google Sheets, Notion ou banco vetorizado)

1. OBJETIVO: HIPERTROFIA (BULKING)
Categoria
Conteúdo
Macros recomendados
2.0–2.5g/kg proteína • 4–6g/kg carbo • 1g/kg gordura (valores médios)
Proteínas preferidas
Frango, patinho, ovos, whey, tilápia, tofu firme, grão-de-bico, lentilha
Carboidratos chave
Arroz branco/integral, batata doce, mandioca, aveia, macarrão integral
Gorduras boas
Azeite, pasta de amendoim, abacate, castanha, óleo de coco
Receitas base
Omelete monstro, Bowl hiperproteico, Arroz bomba com frango
Suplementos sugeridos
Whey, creatina, maltodextrina ou palatinose
Duração média de preparo
15–25 minutos
Dicas adicionais
Fracionar carbo alto ao redor do treino; sempre incluir 1 fonte de gordura boa na refeição principal


2. OBJETIVO: CUTTING (DEFINIÇÃO / SECAR)
Categoria
Conteúdo
Macros recomendados
2.2–2.8g/kg proteína • 2–3g/kg carbo • 0.8g/kg gordura
Proteínas preferidas
Peito de frango, claras, tilápia, tofu firme, proteína vegetal isolada
Carboidratos chave
Abobrinha, couve-flor, arroz de couve-flor, batata doce, arroz integral
Gorduras boas
Azeite (quantidade controlada), linhaça, chia, óleo de coco
Receitas base
Panqueca de claras com aveia, Bowl seco de frango com legumes assados
Suplementos sugeridos
Whey isolado, cafeína, termogênicos (quando aplicável)
Duração média de preparo
10–20 minutos
Dicas adicionais
Evitar combo de gordura + carbo alto na mesma refeição; foco em saciedade e baixo volume calórico


3. OBJETIVO: ENDURANCE (CORRIDA, TRIATLO, ETC.)
Categoria
Conteúdo
Macros recomendados
1.6–2.2g/kg proteína • 5–7g/kg carbo • 1g/kg gordura
Proteínas preferidas
Frango, peixe, ovos, lentilha, tofu
Carboidratos chave
Batata inglesa, arroz branco, frutas secas, aveia, banana, mel, pão
Gorduras boas
Sementes, azeite, pasta de amendoim (quantidade moderada)
Receitas base
Mingau de aveia com whey, Pasta integral com frango e abobrinha
Suplementos sugeridos
Isotônicos, palatinose, BCAA, whey
Duração média de preparo
10–30 minutos (dependendo da refeição pré ou pós treino)
Dicas adicionais
Refeições pré-treino leves e ricas em carbo; priorizar proteína de digestão fácil no pós


Seja simples, direto e prático. Use alimentos comuns no Brasil e foque em performance.
"""

    with st.spinner("Gerando sua receita com IA... 🍳"):
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Você é um nutricionista funcional especializado em nutrição esportiva personalizada."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1000
        )
        resposta = response.choices[0].message.content

    st.markdown("---")
    st.subheader("🍽️ Receita NutriChef Pro:")
    st.markdown(resposta)

    # ========== GERAR RELATÓRIO EM .DOCX ==========
    doc = Document()
    doc.add_heading("Receita NutriChef Pro", level=1)
    for linha in resposta.split("\n"):
        if linha.strip():
            paragrafo = doc.add_paragraph()
            paragrafo.add_run(linha.strip()).font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📥 Baixar receita como .docx",
        data=buffer,
        file_name="Receita_NutriChef_Pro.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # ========== CHAMADA PARA FUNIL ==========
    st.markdown("---")
    st.success("💡 Curtiu a receita?")
    st.markdown("Receba um plano completo de 3 ou 7 dias com IA:")
    col1, col2 = st.columns(2)
    with col1:
        st.link_button("📅 Plano de 3 dias", "https://seulink.com/plano3")
    with col2:
        st.link_button("🔥 Quero o PRO+ 7 dias", "https://seulink.com/assinatura")
